#!/usr/bin/env python3
"""
Generates the real runbook documents (.docx and .pdf) used by the
LocalStack lab evaluation harness. These are genuine Word/PDF documents
with real troubleshooting content -- an agent that calls
`read_runbook_document` gets back text extracted from an actual .docx/.pdf
file on disk, not a hardcoded string baked into Python.

Run:
    python3 data/runbook_docs/generate_docs.py
"""
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt
from fpdf import FPDF

OUT_DIR = Path(__file__).parent

RUNBOOKS = {
    "RB_SCHEMA_DRIFT_DOC": {
        "title": "Runbook: Schema Drift in Customer Profile Ingest",
        "format": "docx",
        "sections": [
            ("Symptom", "The nemoguard-ingest-job Lambda raises a KeyError while validating an "
                        "incoming customer_profile record. CloudWatch Alarm nemoguard-ingest-job-errors "
                        "transitions to ALARM."),
            ("Likely Root Cause", "An upstream producer changed the customer_profile JSON schema "
                                   "without notice -- most commonly a required field (last_login_ip) "
                                   "was removed or renamed in a recent deploy."),
            ("Diagnostic Steps", "1. Call read_s3_object on the exact key referenced in the failing "
                                  "run's logs to confirm which field is actually missing.\n"
                                  "2. Call list_recent_changes against the producer service to check for "
                                  "a correlated recent deployment.\n"
                                  "3. Call query_cloudwatch_logs on /aws/lambda/nemoguard-ingest-job "
                                  "to confirm this is the FIRST occurrence, not a sustained pattern."),
            ("Remediation Steps", "1. Do NOT simply retry -- the same malformed object will fail again.\n"
                                   "2. If this is an isolated bad record: quarantine it (do not delete "
                                   "silently) and escalate to the producing team.\n"
                                   "3. If the schema change is confirmed intentional and permanent: "
                                   "this requires a code change to the ingest job's validation logic, "
                                   "which is OUT OF SCOPE for an automated recovery action -- escalate "
                                   "to the owning engineering team rather than attempting an automated fix."),
            ("Verification", "After a fix is deployed, re-run the exact same object through the job "
                              "(via a healthy-scenario replay) and confirm it succeeds with no KeyError."),
            ("Risk Classification", "MEDIUM. No data loss risk, but requires human judgment on whether "
                                     "the schema change is a mistake or an intentional producer change."),
        ],
    },
    "RB_PARTIAL_WRITE_DOC": {
        "title": "Runbook: Partial Write Crash in Order Events Job",
        "format": "docx",
        "sections": [
            ("Symptom", "The nemoguard-order-events-job Lambda crashes mid-batch. Some rows for the "
                        "run_id are present in order_events, others are missing."),
            ("Likely Root Cause", "The job commits each row in its own transaction rather than one "
                                   "atomic transaction for the whole batch (this matches how many real "
                                   "Spark/Glue JDBC sinks behave). A crash partway through therefore "
                                   "leaves genuinely partial data committed."),
            ("Diagnostic Steps", "1. ALWAYS call check_table_staleness for order_events + the failing "
                                  "run_id BEFORE proposing any rerun. This tells you the actual "
                                  "committed row count vs. the expected count from the run manifest.\n"
                                  "2. If is_stale_or_partial is true, partial data exists and MUST be "
                                  "cleaned up before any rerun."),
            ("Remediation Steps", "1. Call cleanup_partial_write with dry_run=true FIRST and review the "
                                   "row count that would be deleted.\n"
                                   "2. Only after review, call cleanup_partial_write with dry_run=false.\n"
                                   "3. Rerun the job for this run_id.\n"
                                   "4. NEVER rerun before cleanup -- this double-writes the rows that "
                                   "already committed successfully, corrupting downstream aggregates."),
            ("Verification", "Call verify_row_count_matches_expected for order_events + the run_id + "
                              "the manifest's expected_row_count. Only mark the incident RESOLVED if "
                              "this returns verified=true."),
            ("Risk Classification", "MEDIUM. The cleanup step is destructive (DELETEs committed rows) "
                                     "and requires human approval before execution without dry_run."),
        ],
    },
    "RB_POISON_PILL_DOC": {
        "title": "Runbook: Poison-Pill Message in Notification Queue",
        "format": "docx",
        "sections": [
            ("Symptom", "The nemoguard-notification-job Lambda repeatedly fails on the SAME message. "
                        "get_sqs_queue_attributes shows ApproximateNumberOfMessages growing over time "
                        "-- the queue is backing up because one message keeps failing and being "
                        "redelivered, blocking every valid message behind it."),
            ("Likely Root Cause", "A single malformed message (missing a required field, e.g. user_id) "
                                   "cannot be processed by the consumer. Without a dead-letter queue "
                                   "redrive policy, SQS keeps redelivering it after each visibility "
                                   "timeout, and the consumer keeps crashing on it."),
            ("Diagnostic Steps", "1. Call peek_sqs_messages on the queue to inspect message bodies "
                                  "WITHOUT deleting them, and identify which one is malformed.\n"
                                  "2. Call get_sqs_queue_attributes to quantify how backed up the queue "
                                  "actually is (ApproximateNumberOfMessages vs. "
                                  "ApproximateNumberOfMessagesNotVisible)."),
            ("Remediation Steps", "1. The single poison message should be moved aside (e.g. to a "
                                   "dead-letter queue or quarantine location) so it stops blocking the "
                                   "queue -- this is a data-integrity-sensitive action and requires "
                                   "human approval, since it involves deciding the fate of a real "
                                   "customer-facing notification.\n"
                                   "2. Do NOT simply purge the entire queue -- that would silently drop "
                                   "every valid pending notification, not just the bad one."),
            ("Verification", "After the poison message is removed, re-check "
                              "get_sqs_queue_attributes and confirm ApproximateNumberOfMessages is "
                              "decreasing again (queue is draining)."),
            ("Risk Classification", "HIGH. Requires approval -- touches real customer notification "
                                     "delivery, and an incorrect action (purging the whole queue) would "
                                     "cause real data/notification loss."),
        ],
    },
    "RB_PIPELINE_STEP_FAILURE_DOC": {
        "title": "Runbook: Orchestrated Pipeline Step Failure (Step Functions)",
        "format": "pdf",
        "sections": [
            ("Symptom", "The nemoguard-daily-pipeline Step Functions execution enters a FAILED "
                        "status. One of its two states (IngestCustomerProfile or "
                        "ProcessOrderEvents) raised an unhandled exception."),
            ("Likely Root Cause", "The root cause is whichever Lambda the FAILED state actually "
                                   "invoked, NOT the orchestration layer itself. Step Functions is "
                                   "just the wrapper -- diagnose the underlying Lambda's real failure "
                                   "using the same runbooks as its standalone counterpart."),
            ("Diagnostic Steps", "1. Call describe_step_function_execution with the execution ARN to "
                                  "see exactly which state failed and the failure event details.\n"
                                  "2. If IngestCustomerProfile failed: follow RB_SCHEMA_DRIFT_DOC.\n"
                                  "3. If ProcessOrderEvents failed: follow RB_PARTIAL_WRITE_DOC.\n"
                                  "4. Because ProcessOrderEvents only runs AFTER IngestCustomerProfile "
                                  "succeeds, if IngestCustomerProfile failed, ProcessOrderEvents never "
                                  "ran at all -- do not treat this as a partial-write scenario."),
            ("Remediation Steps", "1. Fix the root-cause Lambda per its own runbook.\n"
                                   "2. Re-invoke the ENTIRE state machine execution from the start "
                                   "(StartAt), not just the failed state -- Step Functions standard "
                                   "executions do not support resuming from a specific state.\n"
                                   "3. If IngestCustomerProfile succeeded before ProcessOrderEvents "
                                   "failed, re-running the whole execution will re-invoke "
                                   "IngestCustomerProfile too -- confirm this is idempotent (it is, "
                                   "for this lab's ingest job) before doing so, or split remediation "
                                   "into a manual two-step process if not."),
            ("Verification", "Call describe_step_function_execution on the NEW execution ARN and "
                              "confirm status is SUCCEEDED with no recent_failure_events."),
            ("Risk Classification", "MEDIUM. Re-running an entire multi-step execution can have "
                                     "side effects beyond the failed step if upstream steps are not "
                                     "verified idempotent."),
        ],
    },
}


def _write_docx(runbook_id: str, title: str, sections: list) -> Path:
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Runbook ID: {runbook_id}").italic = True
    for heading, body in sections:
        doc.add_heading(heading, level=2)
        p = doc.add_paragraph(body)
        p.style.font.size = Pt(11)
    out_path = OUT_DIR / f"{runbook_id}.docx"
    doc.save(out_path)
    return out_path


def _write_pdf(runbook_id: str, title: str, sections: list) -> Path:
    epw = 180  # effective page width in mm, well within an A4 page with margins
    pdf = FPDF()
    pdf.set_margins(15, 15, 15)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    def line(text, size, style=""):
        pdf.set_font("Helvetica", style, size)
        pdf.set_x(15)
        pdf.multi_cell(epw, size * 0.6, text)

    line(title, 16, "B")
    line(f"Runbook ID: {runbook_id}", 10, "I")
    pdf.ln(2)
    for heading, body in sections:
        line(heading, 13, "B")
        line(body, 11, "")
        pdf.ln(2)
    out_path = OUT_DIR / f"{runbook_id}.pdf"
    pdf.output(str(out_path))
    return out_path


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for runbook_id, spec in RUNBOOKS.items():
        if spec["format"] == "docx":
            path = _write_docx(runbook_id, spec["title"], spec["sections"])
        else:
            path = _write_pdf(runbook_id, spec["title"], spec["sections"])
        print(f"  Wrote {path}")
    print(f"\nGenerated {len(RUNBOOKS)} runbook documents in {OUT_DIR}")


if __name__ == "__main__":
    main()
